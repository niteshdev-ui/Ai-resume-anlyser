import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import docx
import pdfplumber


@dataclass
class ParsedResume:
    filename: str
    text: str
    name: Optional[str] = None
    email: Optional[str] = None
    phone_number: Optional[str] = None
    skills: List[str] = field(default_factory=list)
    education: List[str] = field(default_factory=list)
    experience: List[str] = field(default_factory=list)
    summary: Optional[str] = None
    ats_score: float = 0.0
    skill_match_percentage: float = 0.0
    resume_strength_score: float = 0.0
    matched_skills: List[str] = field(default_factory=list)
    missing_skills: List[str] = field(default_factory=list)
    keywords_to_add: List[str] = field(default_factory=list)
    missing_sections: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)
    career_recommendations: List[str] = field(default_factory=list)
    keyword_density: float = 0.0
    section_quality_score: float = 0.0
    grammar_score: float = 0.0
    readability_score: float = 0.0

    @property
    def candidate_name(self) -> str:
        return self.name or self.filename


class ResumeParser:
    """Parse PDF/DOCX/TXT resumes and extract useful fields."""

    TECH_SKILLS = {
        "python",
        "java",
        "javaScript",
        "javascript",
        "typescript",
        "c++",
        "c#",
        "sql",
        "postgresql",
        "mysql",
        "mongodb",
        "redis",
        "aws",
        "azure",
        "docker",
        "kubernetes",
        "linux",
        "git",
        "github",
        "ci/cd",
        "pytest",
        "pandas",
        "numpy",
        "machine learning",
        "deep learning",
        "artificial intelligence",
        "streamlit",
        "flask",
        "django",
        "fastapi",
        "spark",
        "hadoop",
        "tableau",
        "power bi",
        "excel",
        "html",
        "css",
        "react",
        "node.js",
        "nodejs",
        "rest api",
        "api",
        "tensorflow",
        "pytorch",
        "opencv",
        "nlp",
        "spacy",
        "pyspark",
        "airflow",
    }

    def parse(self, uploaded_file) -> ParsedResume:
        filename = getattr(uploaded_file, "name", "resume")
        content = uploaded_file.getvalue()
        self._validate_file(filename, content)

        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            text = self._extract_pdf_text(content)
        elif suffix == ".docx":
            text = self._extract_docx_text(content)
        elif suffix in {".txt", ".md"}:
            text = content.decode("utf-8", errors="ignore")
        else:
            raise ValueError("Unsupported file format")

        return self.parse_text(filename, text)

    def parse_text(self, filename: str, text: str) -> ParsedResume:
        cleaned_text = self._normalize_text(text)
        name = self._extract_name(cleaned_text)
        email = self._extract_email(cleaned_text)
        phone_number = self._extract_phone_number(cleaned_text)
        skills = self._extract_skills(cleaned_text)
        education = self._extract_education(cleaned_text)
        experience = self._extract_experience(cleaned_text)
        summary = self._generate_summary(name, skills, education, experience)

        return ParsedResume(
            filename=filename,
            text=cleaned_text,
            name=name,
            email=email,
            phone_number=phone_number,
            skills=skills,
            education=education,
            experience=experience,
            summary=summary,
        )

    def _validate_file(self, filename: str, content: bytes) -> None:
        allowed_extensions = {".pdf", ".docx", ".txt", ".md"}
        suffix = Path(filename).suffix.lower()
        if suffix not in allowed_extensions:
            raise ValueError("Unsupported file extension")
        if len(content) > 10 * 1024 * 1024:
            raise ValueError("File is too large. Maximum size is 10 MB")
        if b"<script" in content.lower() or b"<html" in content.lower():
            raise ValueError("Potentially unsafe file content detected")

    def _extract_pdf_text(self, content: bytes) -> str:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    def _extract_docx_text(self, content: bytes) -> str:
        document = docx.Document(io.BytesIO(content))
        text_parts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        return "\n".join(text_parts)

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _extract_name(self, text: str) -> Optional[str]:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return None
        first_line = lines[0]
        if len(first_line.split()) <= 4 and not re.search(r"@|\d", first_line):
            return first_line
        return None

    def _extract_email(self, text: str) -> Optional[str]:
        match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
        return match.group(0) if match else None

    def _extract_phone_number(self, text: str) -> Optional[str]:
        match = re.search(r"(?:\+?\d[\d -]{8,}\d)", text)
        return match.group(0).strip() if match else None

    def _extract_skills(self, text: str) -> List[str]:
        lower_text = text.lower()
        found_skills = []
        for skill in sorted(self.TECH_SKILLS, key=len, reverse=True):
            pattern = re.escape(skill.lower())
            if re.search(r"\b" + pattern + r"\b", lower_text):
                found_skills.append(skill)
        return found_skills

    def _extract_education(self, text: str) -> List[str]:
        education_patterns = [
            r"bsc",
            r"b\.tech",
            r"m\.tech",
            r"master",
            r"bachelor",
            r"phd",
            r"diploma",
        ]
        matches = []
        for pattern in education_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                matches.append(pattern)
        return matches

    def _extract_experience(self, text: str) -> List[str]:
        experience_lines = []
        for line in text.splitlines():
            normalized = line.strip()
            if re.search(r"years|experience|developer|engineer|analyst|manager|lead", normalized, re.IGNORECASE):
                experience_lines.append(normalized)
        return experience_lines[:8]

    def _generate_summary(self, name, skills, education, experience) -> str:
        summary_parts = []
        if name:
            summary_parts.append(f"{name} brings experience across core technical domains.")
        if skills:
            summary_parts.append(f"Key strengths include {', '.join(skills[:5])}.")
        if education:
            summary_parts.append("The profile shows academic preparation relevant to the target role.")
        if experience:
            summary_parts.append("The candidate has professional experience that aligns with the described position.")
        return " ".join(summary_parts) if summary_parts else "Resume summary not available."
